import sys
import os
import io
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
    QGroupBox, QCheckBox, QLabel, QDoubleSpinBox, QSpinBox, 
    QPushButton, QRadioButton, QScrollArea, QFileDialog, QMessageBox, QGridLayout,
    QStackedWidget, QSizePolicy
)
from PyQt5.QtCore import Qt, QTimer

try:
    import docx
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Safely import all solver modules
try:
    import solvers.finite_difference_Polytropic as fdm_poly
    import solvers.finite_element_Polytropic as fem_poly
    import solvers.spectral_Polytropic as spec_poly
    import solvers.shooting_RK4_Polytropic as rk4_poly
    import solvers.series_ADM_Polytropic as adm_poly
    import solvers.b_spline_Polytropic as bsp_poly
    SOLVERS_AVAILABLE = True
except ImportError:
    try:
        import finite_difference_Polytropic as fdm_poly
        import finite_element_Polytropic as fem_poly
        import spectral_Polytropic as spec_poly
        import shooting_RK4_Polytropic as rk4_poly
        import series_ADM_Polytropic as adm_poly
        import b_spline_Polytropic as bsp_poly
        SOLVERS_AVAILABLE = True
    except ImportError:
        SOLVERS_AVAILABLE = False


class MplCanvas(FigureCanvas):
    """In-memory Matplotlib Canvas for dynamic PyQt embedding"""
    def __init__(self, parent=None, width=8, height=6, dpi=100):
        self.fig = Figure(figsize=(width, height), dpi=dpi)
        super().__init__(self.fig)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.updateGeometry()


class LaneEmdenGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Lane-Emden Numerical Methods Lab")
        self.setGeometry(50, 50, 1400, 950)
        
        self.plot_queue = [] 
        self.export_tables = {}     
        self.current_fig_index = 0
        
        self.timer = QTimer()
        self.timer.timeout.connect(self.display_next_automated_figure)

        self.init_ui()

    def init_ui(self):
        self.setStyleSheet("""
            QWidget { background-color: #121212; color: #E0E0E0; font-family: 'Segoe UI', Arial, sans-serif; }
            QGroupBox { font-weight: bold; border: 1px solid #333; margin-top: 15px; padding-top: 15px; border-radius: 6px; }
            QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 8px; color: #FF5722; font-size: 13px; }
            QPushButton { background-color: #262626; border: 1px solid #444; padding: 10px; border-radius: 5px; font-weight: bold; }
            QPushButton:hover { background-color: #333333; border-color: #666; }
            #RunButton { background-color: #D84315; color: white; font-size: 16px; font-weight: bold; border: none; min-height: 50px; }
            #RunButton:hover { background-color: #F4511E; }
            #NavButton { background-color: #0288D1; color: white; border: none; min-height: 35px; font-size: 14px; }
            #NavButton:hover { background-color: #039BE5; }
            #ExportButton { min-height: 45px; font-size: 13px; }
            QDoubleSpinBox, QSpinBox { background-color: #1E1E1E; border: 1px solid #333; padding: 6px; border-radius: 4px; color: #FFF; }
            QCheckBox { spacing: 8px; }
            QCheckBox::indicator { width: 18px; height: 18px; border-radius: 3px; border: 1px solid #555; background-color: #1E1E1E; }
            QCheckBox::indicator:checked { background-color: #FF5722; border-color: #FF5722; }
        """)

        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.main_widget = QWidget()
        self.main_layout = QVBoxLayout(self.main_widget)
        self.main_layout.setContentsMargins(20, 20, 20, 20)
        self.main_layout.setSpacing(15)

        header = QLabel("Lane-Emden Numerical Methods Lab")
        header.setStyleSheet("font-size: 28px; font-weight: bold; font-family: Georgia, serif; color: #FFFFFF; margin-bottom: 10px;")
        header.setAlignment(Qt.AlignCenter)
        self.main_layout.addWidget(header)

        split_layout = QHBoxLayout()
        self.p1_panel = self.build_problem_panel("Problem 1: Isothermal Gas Sphere", "Lambda (λ)")
        self.p2_panel = self.build_problem_panel("Problem 2: Polytropic Model", "Polytropic Index (n)")
        split_layout.addWidget(self.p1_panel['group'])
        split_layout.addWidget(self.p2_panel['group'])
        self.main_layout.addLayout(split_layout)

        self.main_layout.addWidget(self.build_execution_panel())

        self.gallery_group = QGroupBox("3D Figure Gallery")
        gallery_layout = QVBoxLayout()
        
        self.gallery_stack = QStackedWidget()
        self.gallery_stack.setMinimumHeight(600) 
        
        self.placeholder = QLabel("Configure parameters and click RUN SOLVERS to generate 3D figures.")
        self.placeholder.setAlignment(Qt.AlignCenter)
        self.placeholder.setStyleSheet("border: 2px dashed #444; color: #888; font-size: 16px; border-radius: 8px;")
        self.gallery_stack.addWidget(self.placeholder)
        
        self.canvas = MplCanvas(self, width=8, height=6, dpi=100)
        self.gallery_stack.addWidget(self.canvas)
        
        gallery_layout.addWidget(self.gallery_stack)

        nav_layout = QHBoxLayout()
        self.btn_prev_fig = QPushButton("◄ Previous Figure")
        self.btn_prev_fig.setObjectName("NavButton")
        self.btn_prev_fig.clicked.connect(self.show_previous_figure)
        
        self.lbl_fig_status = QLabel("Figure 0 / 0")
        self.lbl_fig_status.setAlignment(Qt.AlignCenter)
        self.lbl_fig_status.setStyleSheet("font-weight: bold; font-size: 15px; color: #FFF;")

        self.btn_next_fig = QPushButton("Next Figure ►")
        self.btn_next_fig.setObjectName("NavButton")
        self.btn_next_fig.clicked.connect(self.show_next_figure)

        nav_layout.addWidget(self.btn_prev_fig, stretch=1)
        nav_layout.addWidget(self.lbl_fig_status, stretch=2)
        nav_layout.addWidget(self.btn_next_fig, stretch=1)
        
        gallery_layout.addLayout(nav_layout)
        self.gallery_group.setLayout(gallery_layout)
        self.main_layout.addWidget(self.gallery_group, stretch=1)
        
        self.main_layout.addWidget(self.build_export_panel())

        self.scroll_area.setWidget(self.main_widget)
        self.setCentralWidget(self.scroll_area)

    def build_problem_panel(self, title, param_label):
        group = QGroupBox(title)
        layout = QVBoxLayout()
        layout.setSpacing(15)

        enable_cb = QCheckBox("Enabled")
        enable_cb.setChecked(True)
        enable_cb.setStyleSheet("color: #FF5722; font-weight: bold; font-size: 14px;")
        layout.addWidget(enable_cb, alignment=Qt.AlignRight)

        methods_gb = QGroupBox("NUMERICAL METHODS")
        m_layout = QGridLayout()
        m_layout.setSpacing(10)
        select_all = QCheckBox("Select All")
        m_layout.addWidget(select_all, 0, 0, 1, 2)

        method_list = ["FDM", "FEM", "Chebyshev", "Shooting", "ADM", "B-Spline"]
        cb_dict = {}
        for idx, m in enumerate(method_list):
            cb = QCheckBox(m)
            cb_dict[m] = cb
            m_layout.addWidget(cb, (idx // 2) + 1, idx % 2)

        def toggle_all(state):
            for cb in cb_dict.values():
                cb.setChecked(state == Qt.Checked)
        select_all.stateChanged.connect(toggle_all)
        methods_gb.setLayout(m_layout)
        layout.addWidget(methods_gb)

        params_gb = QGroupBox("PARAMETER VARIATION")
        p_layout = QGridLayout()
        p_layout.setSpacing(10)

        p_layout.addWidget(QLabel(f"{param_label} Sequence:"), 0, 0, 1, 3)
        p_layout.addWidget(QLabel("MIN"), 1, 0)
        p_layout.addWidget(QLabel("MAX"), 1, 1)
        p_layout.addWidget(QLabel("SAMPLES"), 1, 2)

        min_sp = QDoubleSpinBox(); min_sp.setValue(1.0)
        max_sp = QDoubleSpinBox(); max_sp.setValue(5.0)
        
        # Updated to request Samples instead of Step
        samples_sp = QSpinBox(); samples_sp.setValue(4); samples_sp.setMinimum(1); samples_sp.setMaximum(1000)
        
        p_layout.addWidget(min_sp, 2, 0)
        p_layout.addWidget(max_sp, 2, 1)
        p_layout.addWidget(samples_sp, 2, 2)

        p_layout.addWidget(QLabel("DOMAIN (XMAX)"), 3, 0)
        p_layout.addWidget(QLabel("TARGET (YTARGET)"), 3, 1)
        xmax_sp = QDoubleSpinBox(); xmax_sp.setValue(2.0); xmax_sp.setMinimum(0.1)
        ytarget_sp = QDoubleSpinBox(); ytarget_sp.setValue(0.0)
        p_layout.addWidget(xmax_sp, 4, 0)
        p_layout.addWidget(ytarget_sp, 4, 1)

        params_gb.setLayout(p_layout)
        layout.addWidget(params_gb)

        term_gb = QGroupBox("TERMINATION CONDITIONS")
        t_layout = QGridLayout()
        t_layout.setSpacing(10)
        
        t_layout.addWidget(QLabel("E_A (APPROX. ERROR)"), 0, 0)
        t_layout.addWidget(QLabel("E_T (TRUE ERROR)"), 0, 1)
        ea_sp = QDoubleSpinBox(); ea_sp.setDecimals(9); ea_sp.setValue(1e-9)
        et_sp = QDoubleSpinBox(); et_sp.setDecimals(9); et_sp.setValue(1e-9)
        t_layout.addWidget(ea_sp, 1, 0)
        t_layout.addWidget(et_sp, 1, 1)

        t_layout.addWidget(QLabel("MAX ITERATIONS"), 2, 0)
        t_layout.addWidget(QLabel("GRID POINTS (N)"), 2, 1)
        max_iter_sp = QSpinBox(); max_iter_sp.setMaximum(10000); max_iter_sp.setValue(40)
        grid_pts_sp = QSpinBox(); grid_pts_sp.setMaximum(10000); grid_pts_sp.setValue(40)
        t_layout.addWidget(max_iter_sp, 3, 0)
        t_layout.addWidget(grid_pts_sp, 3, 1)

        t_layout.addWidget(QLabel("SERIES TERMS (ADM)"), 4, 0)
        series_sp = QSpinBox(); series_sp.setMaximum(500); series_sp.setValue(20)
        t_layout.addWidget(series_sp, 5, 0)

        term_gb.setLayout(t_layout)
        layout.addWidget(term_gb)
        group.setLayout(layout)
        
        return {
            'group': group, 'enable': enable_cb, 'methods': cb_dict,
            'min': min_sp, 'max': max_sp, 'samples': samples_sp,
            'xmax': xmax_sp, 'ytarget': ytarget_sp,
            'ea': ea_sp, 'et': et_sp, 'max_iter': max_iter_sp,
            'grid_pts': grid_pts_sp, 'series_terms': series_sp
        }

    def build_execution_panel(self):
        panel = QWidget()
        layout = QHBoxLayout(panel)
        layout.setContentsMargins(0, 10, 0, 10)

        modes_gb = QGroupBox("EXECUTION MODE")
        m_layout = QHBoxLayout()
        m_layout.setSpacing(20)

        self.radio_oneshot = QRadioButton("One-Shot")
        self.radio_oneshot.setChecked(True)

        self.radio_delay = QRadioButton("Automated Delay")
        self.delay_spin = QDoubleSpinBox()
        self.delay_spin.setSuffix(" sec")
        self.delay_spin.setValue(1.5)
        self.delay_spin.setSingleStep(0.5)

        self.radio_manual = QRadioButton("Manual / Step-by-Step")

        m_layout.addWidget(self.radio_oneshot)
        m_layout.addWidget(self.radio_delay)
        m_layout.addWidget(self.delay_spin)
        m_layout.addWidget(self.radio_manual)
        modes_gb.setLayout(m_layout)

        self.btn_run = QPushButton("RUN SOLVERS")
        self.btn_run.setObjectName("RunButton")
        self.btn_run.clicked.connect(self.execute_solvers)

        layout.addWidget(modes_gb, stretch=2)
        layout.addWidget(self.btn_run, stretch=1)
        return panel

    def build_export_panel(self):
        panel = QGroupBox("Export Center")
        layout = QHBoxLayout()
        layout.setSpacing(15)

        btn_folder = QPushButton("Save Figures to Folder\n(PNG Images)")
        btn_folder.setObjectName("ExportButton")
        btn_word = QPushButton("Export to Word Document\n(.docx Report)")
        btn_word.setObjectName("ExportButton")
        btn_excel = QPushButton("Export to Excel Sheet\n(.xlsx Worksheets)")
        btn_excel.setObjectName("ExportButton")

        btn_folder.clicked.connect(self.export_figures_to_folder)
        btn_word.clicked.connect(self.export_to_word)
        btn_excel.clicked.connect(self.export_to_excel)

        layout.addWidget(btn_folder)
        layout.addWidget(btn_word)
        layout.addWidget(btn_excel)
        panel.setLayout(layout)
        return panel

    def execute_solvers(self):
        if not SOLVERS_AVAILABLE:
            QMessageBox.critical(self, "Import Error", "Could not import solver scripts. Ensure they are in the correct directory.")
            return

        self.timer.stop()
        self.plot_queue.clear()
        self.export_tables.clear()
        self.gallery_stack.setCurrentWidget(self.placeholder)

        if self.p2_panel['enable'].isChecked():
            p_min = self.p2_panel['min'].value()
            p_max = self.p2_panel['max'].value()
            p_samples = self.p2_panel['samples'].value()
            
            # Apply Professor's Exact Modification Formula
            if p_samples > 0:
                p_step = (p_max - p_min) / p_samples
                param_vals = np.arange(p_min, p_max + (p_step/2), p_step)
            else:
                param_vals = np.array([p_min])
            
            xmax = self.p2_panel['xmax'].value()
            max_iter = self.p2_panel['max_iter'].value()
            n_elem = self.p2_panel['grid_pts'].value()
            series_terms = self.p2_panel['series_terms'].value()
            methods = self.p2_panel['methods']

            for method_name, cb in methods.items():
                if cb.isChecked():
                    if method_name == "FDM" and hasattr(fdm_poly, 'solve_polytropic_fdm'):
                        self.run_fdm_solver(param_vals, xmax, max_iter, n_elem)
                    elif method_name == "FEM" and hasattr(fem_poly, 'solve_polytropic_fem'):
                        self.run_fem_solver(param_vals, xmax, max_iter, n_elem)
                    elif method_name == "Chebyshev" and hasattr(spec_poly, 'solve_polytropic_spectral'):
                        self.run_spectral_solver(param_vals, xmax, max_iter, n_elem)
                    elif method_name == "Shooting" and hasattr(rk4_poly, 'solve_polytropic_rk4'):
                        self.run_rk4_solver(param_vals, xmax, max_iter, n_elem)
                    elif method_name == "ADM" and hasattr(adm_poly, 'solve_polytropic_adm'):
                        self.run_adm_solver(param_vals, xmax, max_iter, n_elem, series_terms)
                    elif method_name == "B-Spline" and hasattr(bsp_poly, 'solve_polytropic_bspline'):
                        self.run_bspline_solver(param_vals, xmax, max_iter, n_elem)

        if not self.plot_queue:
            QMessageBox.information(self, "No Output", "No active methods or problems were selected.")
            return

        self.current_fig_index = 0

        # Run Modes Logic
        if self.radio_oneshot.isChecked():
            self.render_figure_at_index(len(self.plot_queue) - 1)
        elif self.radio_delay.isChecked():
            self.render_figure_at_index(0)
            delay_ms = int(self.delay_spin.value() * 1000)
            self.timer.start(delay_ms)
        elif self.radio_manual.isChecked():
            self.render_figure_at_index(0)

    # -------------------------------------------------------------------------
    # In-Memory Plotting Engine 
    # -------------------------------------------------------------------------
    def queue_surface_plot(self, X, Y, Z, xlabel, ylabel, zlabel, title, cmap, int_y=False):
        """Stores plot configuration in memory instead of saving to disk"""
        self.plot_queue.append({
            'X': X, 'Y': Y, 'Z': Z,
            'xlabel': xlabel, 'ylabel': ylabel, 'zlabel': zlabel,
            'title': title, 'cmap': cmap, 'int_y': int_y
        })

    def _plot_standard_results(self, method_name, module, param_vals, results):
        nodes_common = results[param_vals[0]]['nodes']
        Xg, Mg = np.meshgrid(nodes_common, param_vals)
        
        Zfinal = np.array([results[m]['Y'] for m in param_vals])
        self.queue_surface_plot(
            Xg, Mg, Zfinal, 
            'x value', 'Parameter (m)', 'Final Nodal Values', 
            f'{method_name} Polytropic - Final Nodal Values', 'viridis'
        )

        Zerr_hi = np.array([module.rel_err(results[m]['Y'], results[m]['ref_hi']) for m in param_vals])
        self.queue_surface_plot(
            Xg, Mg, Zerr_hi, 
            'x value', 'Parameter (m)', 'Relative True Error', 
            f'{method_name} Polytropic - Final Relative True Error', 'plasma'
        )

        # Plot EVERY iteration as requested
        for idx, m in enumerate(param_vals):
            history_nodes = results[m].get('history_nodes', [])
            if not history_nodes:
                continue 
                
            ref_hi = results[m]['ref_hi']
            iters = np.arange(len(history_nodes))
            Xi, Ii = np.meshgrid(nodes_common, iters)
            
            Zval = np.array(history_nodes)
            Zerr = np.array([module.rel_err(hh, ref_hi) for hh in history_nodes])

            self.queue_surface_plot(
                Xi, Ii, Zval, 
                'x value', 'Iteration Number', 'Iterative Nodal Values', 
                f'{method_name} Polytropic - Iterative Nodal Values (m={m})', 'viridis', int_y=True
            )
            self.queue_surface_plot(
                Xi, Ii, Zerr, 
                'x value', 'Iteration Number', 'Relative True Error', 
                f'{method_name} Polytropic - Iterative Relative Error (m={m})', 'plasma', int_y=True
            )

        df = pd.DataFrame(Zfinal, index=[f"m={m:.2f}" for m in param_vals], columns=[f"x={x:.2f}" for x in nodes_common])
        self.export_tables[f"{method_name}_Final"] = df

    # -------------------------------------------------------------------------
    # Solver Interfaces
    # -------------------------------------------------------------------------
    def run_fdm_solver(self, param_vals, xmax, max_iter, n_elem):
        res = {}
        for m in param_vals:
            nodes, c, history, nit, conv = fdm_poly.solve_polytropic_fdm(m, n_elem=n_elem, R=xmax, max_iter=max_iter)
            ref_hi = fdm_poly.get_reference_solution(m, nodes)
            res[m] = {'nodes': nodes, 'Y': c, 'history_nodes': history, 'ref_hi': ref_hi}
        self._plot_standard_results("FDM", fdm_poly, param_vals, res)

    def run_fem_solver(self, param_vals, xmax, max_iter, n_elem):
        res = {}
        for m in param_vals:
            nodes, c, history, nit = fem_poly.solve_polytropic_fem(m, n_elem=n_elem, R=xmax, max_iter=max_iter)
            ref_hi = fem_poly.get_true_reference(m, nodes)
            res[m] = {'nodes': nodes, 'Y': c, 'history_nodes': history, 'ref_hi': ref_hi}
        self._plot_standard_results("FEM", fem_poly, param_vals, res)

    def run_spectral_solver(self, param_vals, xmax, max_iter, n_elem):
        res = {}
        for m in param_vals:
            nodes, c, history = spec_poly.solve_polytropic_spectral(m, N=n_elem, R=xmax, max_iter=max_iter)
            ref_hi = spec_poly.get_true_reference(m, nodes)
            res[m] = {'nodes': nodes, 'Y': c, 'history_nodes': history, 'ref_hi': ref_hi}
        self._plot_standard_results("Chebyshev", spec_poly, param_vals, res)

    def run_rk4_solver(self, param_vals, xmax, max_iter, n_elem):
        res = {}
        for m in param_vals:
            nodes, c, history, step_history, d_sol, nit = rk4_poly.solve_polytropic_rk4(m, n_nodes=n_elem, R=xmax, max_iter=max_iter)
            ref_hi = rk4_poly.get_true_reference(m, nodes)
            res[m] = {'nodes': nodes, 'Y': c, 'history_nodes': history, 'ref_hi': ref_hi}
        self._plot_standard_results("Shooting", rk4_poly, param_vals, res)

    def run_adm_solver(self, param_vals, xmax, max_iter, n_elem, series_terms):
        res = {}
        for m in param_vals:
            _, nodes, c, history, o_hist, pade, nit = adm_poly.solve_polytropic_adm(m, n_nodes=n_elem, R=xmax, max_iter=max_iter, N_series=series_terms)
            ref_hi = adm_poly.get_true_reference(m, nodes)
            res[m] = {'nodes': nodes, 'Y': c, 'history_nodes': history, 'ref_hi': ref_hi}
        self._plot_standard_results("ADM", adm_poly, param_vals, res)

    def run_bspline_solver(self, param_vals, xmax, max_iter, n_elem):
        res = {}
        _, _, tau_m, M0_m, _, _ = bsp_poly.build_bspline_operator(3, n_elem, xmax)
        for m in param_vals:
            t, tau, Y_final, history, c_spline, it = bsp_poly.solve_polytropic_bspline(m, p=3, n_elem=n_elem, R=xmax, max_iter=max_iter)
            ref_hi = bsp_poly.get_true_reference(m, tau)
            history_nodes = [M0_m @ h for h in history]
            res[m] = {'nodes': tau, 'Y': Y_final, 'history_nodes': history_nodes, 'ref_hi': ref_hi}
        self._plot_standard_results("B-Spline", bsp_poly, param_vals, res)

    # -------------------------------------------------------------------------
    # On-The-Fly Matplotlib Rendering 
    # -------------------------------------------------------------------------
    def render_figure_at_index(self, index):
        if 0 <= index < len(self.plot_queue):
            data = self.plot_queue[index]
            
            self.canvas.fig.clf()
            ax = self.canvas.fig.add_subplot(111, projection='3d')
            
            surf = ax.plot_surface(data['X'], data['Y'], data['Z'], cmap=data['cmap'], edgecolor='k', alpha=0.9)
            
            ax.set_xlabel(data['xlabel'])
            ax.set_ylabel(data['ylabel'])
            ax.set_zlabel(data['zlabel'])
            ax.set_title(data['title'], pad=10)
            
            if data['int_y']:
                ax.yaxis.set_major_locator(MaxNLocator(integer=True))
                
            self.canvas.fig.colorbar(surf, ax=ax, shrink=0.5, pad=0.1)
            self.canvas.draw()
            
            self.gallery_stack.setCurrentWidget(self.canvas)
            self.lbl_fig_status.setText(f"Figure {index + 1} / {len(self.plot_queue)}: {data['title']}")

    def display_next_automated_figure(self):
        if self.current_fig_index < len(self.plot_queue) - 1:
            self.current_fig_index += 1
            self.render_figure_at_index(self.current_fig_index)
        else:
            self.timer.stop()

    def show_previous_figure(self):
        if self.current_fig_index > 0:
            self.current_fig_index -= 1
            self.render_figure_at_index(self.current_fig_index)

    def show_next_figure(self):
        if self.current_fig_index < len(self.plot_queue) - 1:
            self.current_fig_index += 1
            self.render_figure_at_index(self.current_fig_index)

    # -------------------------------------------------------------------------
    # Export Handlers 
    # -------------------------------------------------------------------------
    def export_figures_to_folder(self):
        if not self.plot_queue:
            QMessageBox.warning(self, "Export Error", "No figures available to export.")
            return
        folder = QFileDialog.getExistingDirectory(self, "Select Output Directory")
        if folder:
            for idx, data in enumerate(self.plot_queue):
                fig = plt.figure(figsize=(9, 7))
                ax = fig.add_subplot(111, projection='3d')
                surf = ax.plot_surface(data['X'], data['Y'], data['Z'], cmap=data['cmap'], edgecolor='k', alpha=0.9)
                ax.set_xlabel(data['xlabel']); ax.set_ylabel(data['ylabel']); ax.set_zlabel(data['zlabel'])
                ax.set_title(data['title'], pad=10)
                if data['int_y']: ax.yaxis.set_major_locator(MaxNLocator(integer=True))
                fig.colorbar(surf, ax=ax, shrink=0.5, pad=0.1)
                
                clean_title = data['title'].replace(" ", "_").replace("-", "_").replace(":", "")
                dest = os.path.join(folder, f"{idx+1:02d}_{clean_title}.png")
                fig.savefig(dest, bbox_inches='tight')
                plt.close(fig)
            QMessageBox.information(self, "Success", f"Exported {len(self.plot_queue)} images to {folder}")

    def export_to_word(self):
        if not DOCX_AVAILABLE:
            QMessageBox.critical(self, "Error", "python-docx is not installed.")
            return
        if not self.plot_queue:
            return
        filepath, _ = QFileDialog.getSaveFileName(self, "Save Word Document", "", "Word Document (*.docx)")
        if filepath:
            doc = docx.Document()
            doc.add_heading("Numerical Methods: 3D Surface Analysis", level=0)
            
            for idx, data in enumerate(self.plot_queue):
                doc.add_heading(f"Figure {idx+1}: {data['title']}", level=2)
                
                fig = plt.figure(figsize=(8, 6))
                ax = fig.add_subplot(111, projection='3d')
                surf = ax.plot_surface(data['X'], data['Y'], data['Z'], cmap=data['cmap'], edgecolor='k', alpha=0.9)
                ax.set_xlabel(data['xlabel']); ax.set_ylabel(data['ylabel']); ax.set_zlabel(data['zlabel'])
                ax.set_title(data['title'], pad=10)
                if data['int_y']: ax.yaxis.set_major_locator(MaxNLocator(integer=True))
                fig.colorbar(surf, ax=ax, shrink=0.5, pad=0.1)
                
                img_stream = io.BytesIO()
                fig.savefig(img_stream, format='png', bbox_inches='tight')
                plt.close(fig)
                img_stream.seek(0)
                
                doc.add_picture(img_stream, width=Inches(6.0))
                doc.add_paragraph(f"Generated solution data for {data['title']}.")
                
            doc.save(filepath)
            QMessageBox.information(self, "Success", f"Report saved at {filepath}")

    def export_to_excel(self):
        if not EXCEL_AVAILABLE:
            QMessageBox.critical(self, "Error", "openpyxl is not installed.")
            return
        if not self.export_tables:
            return
        filepath, _ = QFileDialog.getSaveFileName(self, "Save Excel Workbook", "", "Excel Workbook (*.xlsx)")
        if filepath:
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                for sheet_name, df in self.export_tables.items():
                    df.to_excel(writer, sheet_name=sheet_name[:31])
            QMessageBox.information(self, "Success", f"Workbook saved at {filepath}")


if __name__ == '__main__':
    app = QApplication(sys.argv)
    gui = LaneEmdenGUI()
    gui.show()
    sys.exit(app.exec_())